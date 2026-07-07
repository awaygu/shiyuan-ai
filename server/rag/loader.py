"""Document loader: parse PDF, DOCX, TXT, MD, Image files with page tracking."""

from __future__ import annotations

import base64
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

# 清洗 PDF 解析乱码：cid(...) 占位符、Unicode 替换字符、控制字符、零宽字符等
_GARBLED_RE = re.compile(
    r"cid\(\d+\)"  # pdfplumber 的 cid 占位符
    r"|[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]"  # C0/C1 控制字符（保留 \n\t\r）
    r"|[​-‏-﻿]"  # 零宽字符、换行分隔符、BOM
    r"|�"  # Unicode 替换字符（�）
)


def _clean_text(text: str) -> str:
    """清洗文本中的乱码字符和不可见字符，保留有意义的空白。"""
    if not text:
        return ""
    text = _GARBLED_RE.sub("", text)
    # 合并连续空白
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


@dataclass
class PageText:
    page: int
    text: str


@dataclass
class Document:
    doc_id: str = field(default_factory=lambda: "")
    filename: str = ""
    text: str = ""
    file_type: str = ""
    file_size: int = 0
    metadata: dict = field(default_factory=dict)
    pages: list[PageText] = field(default_factory=list)
    generated_name: str = ""


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"} | IMAGE_EXTENSIONS

    def load(self, file_path: Path, doc_id: str = "") -> Document:
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        file_size = file_path.stat().st_size
        generated_name = ""

        if ext == ".pdf":
            pages = self._load_pdf(file_path)
            text = "\n\n".join(p.text for p in pages)
        elif ext in (".docx", ".doc"):
            pages = self._load_docx(file_path)
            text = "\n\n".join(p.text for p in pages)
        elif ext in IMAGE_EXTENSIONS:
            pages, generated_name = self._load_image(file_path)
            text = "\n\n".join(p.text for p in pages)
        else:
            pages = self._load_text(file_path)
            text = "\n\n".join(p.text for p in pages)

        # 清洗所有文本，去除 PDF 乱码和不可见字符
        text = _clean_text(text)
        for p in pages:
            p.text = _clean_text(p.text)

        return Document(
            doc_id=doc_id,
            filename=file_path.name,
            text=text,
            file_type=ext,
            file_size=file_size,
            pages=pages,
            generated_name=generated_name,
        )

    def _load_pdf(self, path: Path) -> list[PageText]:
        import pdfplumber

        pages: list[PageText] = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    pages.append(PageText(page=i, text=page_text))
        return pages

    def _load_docx(self, path: Path) -> list[PageText]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        per_page = max(1, len(paragraphs) // max(1, len(paragraphs) // 20))
        pages: list[PageText] = []
        page_num = 1
        for i in range(0, len(paragraphs), per_page):
            chunk = "\n\n".join(paragraphs[i : i + per_page])
            pages.append(PageText(page=page_num, text=chunk))
            page_num += 1
        if not pages and full_text.strip():
            pages.append(PageText(page=1, text=full_text))
        return pages

    def _load_image(self, path: Path) -> list[PageText]:
        from io import BytesIO

        from PIL import Image
        from langchain_core.messages import HumanMessage, SystemMessage
        from langchain_openai import ChatOpenAI

        from config import DASHSCOPE_API_KEY, KB_VISION_BASE_URL, KB_VISION_MODEL, LLM_API_KEY, LLM_BASE_URL, LLM_MODEL
        from core.style_manager import prompt_manager

        img = Image.open(str(path))
        if img.mode == "RGBA" or img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        buf = BytesIO()
        fmt = "PNG" if path.suffix.lower() in (".png", ".webp", ".bmp") else "JPEG"
        img.save(buf, format=fmt)
        b64 = base64.b64encode(buf.getvalue()).decode()
        mime = "image/png" if fmt == "PNG" else "image/jpeg"
        data_url = f"data:{mime};base64,{b64}"

        vision_llm = ChatOpenAI(
            api_key=DASHSCOPE_API_KEY,
            base_url=KB_VISION_BASE_URL,
            model=KB_VISION_MODEL,
            timeout=120,
            max_retries=2,
        )
        ocr_msg = HumanMessage(
            content=[
                {"type": "image_url", "image_url": {"url": data_url}},
                {"type": "text", "text": prompt_manager.ocr_extraction_prompt},
            ]
        )
        ocr_text = vision_llm.invoke([ocr_msg]).content or ""
        if not ocr_text.strip():
            return [], ""

        title_llm = ChatOpenAI(
            api_key=LLM_API_KEY,
            base_url=LLM_BASE_URL,
            model=LLM_MODEL,
            timeout=30,
            max_retries=2,
        )
        title_text = title_llm.invoke([
            SystemMessage(content=prompt_manager.title_generator_system_prompt),
            HumanMessage(content=ocr_text.strip()),
        ]).content or ""
        summary_name = title_text.strip()[:15]

        return [PageText(page=1, text=ocr_text.strip())], summary_name

    def _load_text(self, path: Path) -> list[PageText]:
        encodings = ["utf-8", "gbk", "gb2312", "gb18030"]
        for enc in encodings:
            try:
                full_text = path.read_text(encoding=enc)
                lines = full_text.splitlines()
                per_page = 50
                pages: list[PageText] = []
                page_num = 1
                for i in range(0, len(lines), per_page):
                    pages.append(PageText(page=page_num, text="\n".join(lines[i : i + per_page])))
                    page_num += 1
                if not pages and full_text.strip():
                    pages.append(PageText(page=1, text=full_text))
                return pages
            except (UnicodeDecodeError, LookupError):
                continue
        raise RuntimeError(f"Cannot decode file: {path}")
